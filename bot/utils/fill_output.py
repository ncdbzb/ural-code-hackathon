from docx import Document

def get_output_docx(dict_1, dict_2, dict_3, dict_4, save_path, comment_list):
    # Список словарей и заголовков для первого столбца
    dicts_and_titles = [
        ('структура', dict_1),
        ('введение', dict_2),
        ('заключение', dict_3),
        ('список литературы', dict_4)
    ]

    # Создание документа
    doc = Document()

    # Добавляем заголовки таблицы
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Пункт'
    hdr_cells[1].text = 'Выполнение требования'
    hdr_cells[2].text = 'Комментарий'  # Можно добавить еще один столбец, если требуется

    # Заполнение таблицы данными
    for title, data_dict in dicts_and_titles:
        for key, value in data_dict.items():
            row_cells = table.add_row().cells
            row_cells[0].text = title  # Первый столбец с названием
            row_cells[1].text = key    # Заполняем пункт из словаря
            row_cells[2].text = value  # Заполняем выполнение требования из словаря

    doc.add_paragraph("Комментарии к работе (выполненные с помощью ИИ)")
    doc.add_paragraph("Комментарии к выявленным ошибкам")

    # Цикл по спискам пунктов и вариантов
    for punkt, variant in comment_list:
        doc.add_paragraph(f"Пункт: {punkt}.")
        doc.add_paragraph(f"Возможный вариант: {variant}")
    # Сохранение документа
    doc.save(save_path)
    return
