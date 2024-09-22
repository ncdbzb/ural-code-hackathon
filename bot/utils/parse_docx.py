import re
from docx import Document

def check_structure(doc_path):
    required_elements = [
        'титульный лист',
        'список исполнителей',
        'реферат',
        'содержание',
        'термины и определения',
        'перечень сокращений и обозначений',
        'введение',
        'заключение',
        'список использованных источников',
        'приложения',
    ]
    result_dict = dict()
    doc = Document(doc_path)
    full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    for i in full_text:
        if i.lower() in required_elements:
            result_dict[i.lower()] = '+'

    for i in required_elements:
        if i not in result_dict.keys():
            result_dict[i] = '-'

    return result_dict


def find_table_of_contents(doc_path):
    # Открыть документ
    doc = Document(doc_path)

    # Преобразовать текст документа в список строк
    full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    # Определить, какие ключевые слова могут указать на раздел содержания
    toc_keywords = ['содержание', 'оглавление']

    # Поиск содержания
    toc_start = -1
    for idx, line in enumerate(full_text):
        # Если находим ключевое слово, то запоминаем индекс начала содержания
        if any(keyword.lower() in line.lower() for keyword in toc_keywords):
            toc_start = idx
            break

    if toc_start == -1:
        raise ValueError("Содержание не найдено")

    # Поиск разделов и страниц в содержании
    toc_dict = {}
    section_pattern = re.compile(r'(.+?)\s+(\d+)$')  # Ищем строки вида: "Название раздела ... 12"
    change_text = ''.join(full_text[toc_start].split('Содержание')[1:])
    full_text[toc_start] = change_text
    for line in full_text[toc_start:]:

        if toc_dict.keys():
            if line.lower() in list(toc_dict.keys()):
                break

        # Ищем соответствие шаблону: текст раздела и страница
        match = section_pattern.match(line)
        if match:
            section_name = match.group(1).strip()  # Название раздела
            page_number = int(match.group(2))  # Номер страницы
            toc_dict[section_name.lower()] = page_number

    # Ключевые разделы, которые нужно проверить
    required_sections = ["Введение", "Заключение", "Список использованных источников", "Приложения"]

    # Список отсутствующих разделов
    missing_sections = [section for section in required_sections if
                        section.lower() not in list(map(lambda x: x.lower(), toc_dict.keys()))]
    found_sections = [section.lower() for section in required_sections if
                      section.lower() in list(map(lambda x: x.lower(), toc_dict.keys()))]

    if missing_sections:
        print(f"Не найдены следующие разделы в содержании: {', '.join(missing_sections)}\n")

    return toc_dict, found_sections, missing_sections


def get_paragraphs_after_toc(doc, toc_dict):
    # Извлекаем текст всех параграфов, убирая пустые
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    # Ищем индекс параграфа "Содержание"
    toc_start_index = None
    for i, para in enumerate(paragraphs):
        if "содержание" in para.lower():
            toc_start_index = i
            break

    # Если "Содержание" найдено, обрезаем список paragraphs
    if toc_start_index is not None:
        # Получаем количество элементов содержания из toc_dict
        toc_length = len(toc_dict)
        # Начинаем список с параграфа, следующего за содержанием
        paragraphs = paragraphs[toc_start_index + toc_length + 3:]
        # print(toc_start_index, toc_length)

    return paragraphs


def extract_section_texts_by_toc(doc_path, toc_dict, found_sections):
    # Открываем документ
    doc = Document(doc_path)

    # Преобразуем текст документа в список параграфов
    paragraphs = get_paragraphs_after_toc(doc, toc_dict)

    # Словарь для хранения текстов разделов
    section_texts = {}

    # Получаем список всех заголовков из toc_dict в том порядке, как они идут в документе
    toc_sections = list(toc_dict.keys())

    # Проходим по каждому разделу из found_sections
    for section in found_sections:
        # Ищем индекс начала раздела
        try:
            start_index = next(idx for idx, para in enumerate(paragraphs) if section.lower() in para.lower())
        except StopIteration:
            continue  # Пропускаем, если раздел не найден

        # Ищем индекс следующего раздела из toc_dict
        current_section_idx = toc_sections.index(section)

        if current_section_idx + 1 < len(toc_sections):
            next_section = toc_sections[current_section_idx + 1]
            try:
                end_index = next(idx for idx, para in enumerate(paragraphs) if next_section.lower() in para.lower())
            except StopIteration:
                end_index = len(paragraphs)  # Если следующая секция не найдена, берем до конца документа
        else:
            end_index = len(paragraphs)  # Если это последний раздел, берем до конца документа

        # Извлекаем текст между найденными индексами
        section_texts[section] = "\n".join(paragraphs[start_index:end_index])

    return section_texts

def get_parsing_result(doc_path):
    toc_dict, found_sections, missing_sections = find_table_of_contents(doc_path)
    section_texts = extract_section_texts_by_toc(doc_path, toc_dict, found_sections)
    return section_texts



